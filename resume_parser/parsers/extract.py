from pathlib import Path

from .. import config  # noqa: F401 — Tesseract path from .env when using OCR

import docx
import PyPDF2
from PIL import Image
import pytesseract


def parse_pdf(pdf_path: str) -> str:
    with open(pdf_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        ).strip()


def parse_docx(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs).strip()


def parse_image(image_path: str) -> str:
    image = Image.open(image_path)
    try:
        return pytesseract.image_to_string(image).strip()
    except pytesseract.pytesseract.TesseractNotFoundError as exc:
        hint = (
            "Install Tesseract, add it to PATH, or set TESSERACT_CMD in .env to the full path "
            "of tesseract.exe with no quotes, e.g. TESSERACT_CMD=C:/Program Files/Tesseract-OCR/tesseract.exe "
            "(forward slashes avoid .env backslash issues on Windows)."
        )
        raise ValueError(
            "Tesseract OCR was not found or TESSERACT_CMD does not point to a real file. "
            + hint
        ) from exc


def extract_text_from_file(file_path: str) -> str:
    """Route to the correct parser based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    if ext == ".docx":
        return parse_docx(file_path)
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".webp"):
        return parse_image(file_path)
    if ext == ".txt":
        return Path(file_path).read_text(errors="replace").strip()
    raise ValueError(
        f"Unsupported file type: '{ext}'. Supported: pdf, docx, txt, png, jpg, jpeg, tiff, webp"
    )
